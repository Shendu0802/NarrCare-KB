"""Document parser supporting PDF, Markdown, Docx, JSONL, and plain text."""
import os
from dataclasses import dataclass, field


@dataclass
class Page:
    """A single page of extracted text."""
    page_number: int
    text: str
    parse_method: str = "text_layer"  # text_layer | ocr | hybrid
    confidence: float = 1.0


@dataclass
class ParsedDocument:
    """Result of parsing a document."""
    file_path: str
    source_type: str
    pages: list[Page] = field(default_factory=list)


class DocumentParser:
    """Multi-format document parser.

    Supports: PDF (via PyMuPDF + PaddleOCR fallback), Markdown, Docx, JSONL, TXT.
    """

    TEXT_LAYER_MIN_CHARS = 50

    EXTENSION_MAP = {
        ".pdf": "pdf_book",
        ".md": "markdown",
        ".docx": "markdown",
        ".jsonl": "case",
        ".txt": "markdown",
    }

    def __init__(self, enable_ocr: bool = True):
        self.enable_ocr = enable_ocr
        self._ocr = None

    def _get_ocr(self):
        """Lazy-init OCR engine (easyocr)."""
        if self._ocr is None and self.enable_ocr:
            try:
                import easyocr
                self._ocr = easyocr.Reader(['ch_sim', 'en'], gpu=False, verbose=False)
            except Exception:
                self.enable_ocr = False
        return self._ocr

    def detect_source_type(self, file_path: str) -> str:
        """Infer source_type from file extension."""
        ext = os.path.splitext(file_path)[1].lower()
        return self.EXTENSION_MAP.get(ext, "pdf_book")

    def parse(self, file_path: str, source_type: str | None = None) -> ParsedDocument:
        """Parse a file into a ParsedDocument. Dispatches based on extension."""
        if source_type is None:
            source_type = self.detect_source_type(file_path)

        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".pdf":
            return self._parse_pdf(file_path, source_type)
        elif ext in (".md", ".txt"):
            return self._parse_text(file_path, source_type)
        elif ext == ".docx":
            return self._parse_docx(file_path, source_type)
        elif ext == ".jsonl":
            return self._parse_jsonl(file_path, source_type)
        else:
            return self._parse_text(file_path, source_type)

    def _parse_pdf(self, file_path: str, source_type: str) -> ParsedDocument:
        """Parse PDF using PyMuPDF, with PaddleOCR fallback for scanned pages."""
        import fitz  # PyMuPDF

        doc = fitz.open(file_path)
        pages = []
        ocr_needed = 0

        for i, page in enumerate(doc):
            text = page.get_text()
            if len(text.strip()) >= self.TEXT_LAYER_MIN_CHARS:
                pages.append(Page(
                    page_number=i + 1, text=text.strip(),
                    parse_method="text_layer", confidence=1.0,
                ))
            else:
                ocr_needed += 1
                if self.enable_ocr:
                    ocr_text = self._ocr_page_from_pdf(doc, i)
                    pages.append(Page(
                        page_number=i + 1, text=ocr_text,
                        parse_method="ocr", confidence=0.7,
                    ))
                else:
                    pages.append(Page(
                        page_number=i + 1, text=text.strip(),
                        parse_method="text_layer", confidence=0.3,
                    ))

        doc.close()

        if ocr_needed > 0:
            print(f"  [parser] {file_path}: {ocr_needed}/{len(pages)} pages required OCR")

        return ParsedDocument(file_path=file_path, source_type=source_type, pages=pages)

    def _ocr_page_from_pdf(self, doc, page_index: int) -> str:
        """Extract text from a single PDF page using OCR (easyocr)."""
        import numpy as np
        ocr = self._get_ocr()
        if ocr is None:
            return ""

        page = doc[page_index]
        pix = page.get_pixmap(dpi=120)
        # Convert to numpy array (height, width, channels) RGB
        img_array = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)

        try:
            results = ocr.readtext(img_array)
        except Exception:
            return ""

        lines = []
        for bbox, text, confidence in results:
            if text.strip():
                lines.append(text)
        return "\n".join(lines)

    def _parse_text(self, file_path: str, source_type: str) -> ParsedDocument:
        """Parse plain text or Markdown files."""
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        return ParsedDocument(
            file_path=file_path, source_type=source_type,
            pages=[Page(page_number=1, text=text, parse_method="text_layer", confidence=1.0)],
        )

    def _parse_docx(self, file_path: str, source_type: str) -> ParsedDocument:
        """Parse Docx files."""
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n\n".join(paragraphs)
        return ParsedDocument(
            file_path=file_path, source_type=source_type,
            pages=[Page(page_number=1, text=text, parse_method="text_layer", confidence=1.0)],
        )

    def _parse_jsonl(self, file_path: str, source_type: str) -> ParsedDocument:
        """Parse JSONL files — each line maps to a page."""
        import json
        pages = []
        with open(file_path, "r", encoding="utf-8") as f:
            for i, line in enumerate(f):
                line = line.strip()
                if not line:
                    continue
                try:
                    obj = json.loads(line)
                    text = obj.get("text", obj.get("content", json.dumps(obj, ensure_ascii=False)))
                except json.JSONDecodeError:
                    text = line
                pages.append(Page(page_number=i + 1, text=text, parse_method="text_layer", confidence=1.0))
        return ParsedDocument(file_path=file_path, source_type=source_type, pages=pages)

    def ocr_page(self, file_path: str, page_number: int) -> Page:
        """Run PaddleOCR on a specific page of a PDF. Used as fallback for scanned pages."""
        import fitz
        doc = fitz.open(file_path)
        text = self._ocr_page_from_pdf(doc, page_number - 1)
        doc.close()
        return Page(page_number=page_number, text=text, parse_method="ocr", confidence=0.7)
